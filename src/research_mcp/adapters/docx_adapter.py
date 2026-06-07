"""Local Word/docx adapter built on python-docx."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from research_mcp.adapters import AdapterMeta, BaseAdapter, ToolSpec, register_adapter


@register_adapter
class DocxAdapter(BaseAdapter):
    """Create and edit .docx documents without an external MCP server."""

    adapter_name = "docx"

    def metadata(self) -> AdapterMeta:
        return AdapterMeta(
            name="docx",
            description="Local Word .docx creation and reading tools",
            tools=[
                ToolSpec(
                    name="docx_create",
                    description="Create a .docx document with optional title and paragraphs.",
                    input_schema={
                        "type": "object",
                        "properties": {
                            "output_path": {
                                "type": "string",
                                "description": "Output .docx path",
                                "minLength": 1,
                            },
                            "title": {"type": "string", "description": "Optional document title"},
                            "paragraphs": {
                                "type": "array",
                                "items": {"type": "string"},
                                "description": "Optional initial paragraphs",
                            },
                        },
                        "required": ["output_path"],
                    },
                    handler=self.create,
                ),
                ToolSpec(
                    name="docx_read",
                    description="Read paragraph and table text from a .docx document.",
                    input_schema={
                        "type": "object",
                        "properties": {
                            "docx_path": {
                                "type": "string",
                                "description": "Input .docx path",
                                "minLength": 1,
                            },
                        },
                        "required": ["docx_path"],
                    },
                    handler=self.read,
                ),
                ToolSpec(
                    name="docx_add_heading",
                    description="Append a heading to an existing .docx document.",
                    input_schema={
                        "type": "object",
                        "properties": {
                            "docx_path": {
                                "type": "string",
                                "description": "Input .docx path",
                                "minLength": 1,
                            },
                            "text": {
                                "type": "string",
                                "description": "Heading text",
                                "minLength": 1,
                            },
                            "level": {
                                "type": "integer",
                                "description": "Heading level",
                                "default": 1,
                            },
                        },
                        "required": ["docx_path", "text"],
                    },
                    handler=self.add_heading,
                ),
                ToolSpec(
                    name="docx_add_paragraph",
                    description="Append a paragraph to an existing .docx document.",
                    input_schema={
                        "type": "object",
                        "properties": {
                            "docx_path": {
                                "type": "string",
                                "description": "Input .docx path",
                                "minLength": 1,
                            },
                            "text": {
                                "type": "string",
                                "description": "Paragraph text",
                                "minLength": 1,
                            },
                        },
                        "required": ["docx_path", "text"],
                    },
                    handler=self.add_paragraph,
                ),
                ToolSpec(
                    name="docx_add_table",
                    description="Append a table to an existing .docx document.",
                    input_schema={
                        "type": "object",
                        "properties": {
                            "docx_path": {
                                "type": "string",
                                "description": "Input .docx path",
                                "minLength": 1,
                            },
                            "rows": {
                                "type": "array",
                                "items": {
                                    "type": "array",
                                    "items": {"type": "string"},
                                    "minItems": 1,
                                },
                                "description": "Table rows as string cells",
                                "minItems": 1,
                            },
                        },
                        "required": ["docx_path", "rows"],
                    },
                    handler=self.add_table,
                ),
            ],
        )

    async def initialize(self, config: dict[str, Any] | None = None) -> None:
        pass

    async def create(
        self,
        output_path: str,
        title: str = "",
        paragraphs: list[str] | None = None,
    ) -> dict[str, Any]:
        from docx import Document

        output = Path(output_path).expanduser().resolve()
        output.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        if title:
            document.add_heading(title, level=0)
        for paragraph in paragraphs or []:
            document.add_paragraph(paragraph)
        document.save(output)
        return {"output_path": str(output), "paragraphs": len(paragraphs or [])}

    async def read(self, docx_path: str) -> dict[str, Any]:
        from docx import Document

        path = self._require_docx(docx_path)
        document = Document(path)
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        tables = [
            [[cell.text for cell in row.cells] for row in table.rows] for table in document.tables
        ]
        return {
            "docx_path": str(path),
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
            "paragraphs": paragraphs,
            "tables": tables,
        }

    async def add_heading(self, docx_path: str, text: str, level: int = 1) -> dict[str, Any]:
        from docx import Document

        path = self._require_docx(docx_path)
        document = Document(path)
        document.add_heading(text, level=max(0, min(level, 9)))
        document.save(path)
        return {"docx_path": str(path), "added": "heading", "text": text, "level": level}

    async def add_paragraph(self, docx_path: str, text: str) -> dict[str, Any]:
        from docx import Document

        path = self._require_docx(docx_path)
        document = Document(path)
        document.add_paragraph(text)
        document.save(path)
        return {"docx_path": str(path), "added": "paragraph", "text": text}

    async def add_table(self, docx_path: str, rows: list[list[str]]) -> dict[str, Any]:
        from docx import Document

        if not rows or not rows[0]:
            raise ValueError("rows must contain at least one row and one column")
        width = len(rows[0])
        if any(len(row) != width for row in rows):
            raise ValueError("all rows must have the same number of cells")
        path = self._require_docx(docx_path)
        document = Document(path)
        table = document.add_table(rows=len(rows), cols=width)
        for row_index, row in enumerate(rows):
            for col_index, value in enumerate(row):
                table.cell(row_index, col_index).text = str(value)
        document.save(path)
        return {"docx_path": str(path), "added": "table", "rows": len(rows), "columns": width}

    def _require_docx(self, docx_path: str) -> Path:
        path = Path(docx_path).expanduser().resolve()
        if not path.exists() or not path.is_file():
            raise FileNotFoundError(f"DOCX file not found: {path}")
        if path.suffix.lower() != ".docx":
            raise ValueError(f"Expected a .docx file: {path}")
        return path
